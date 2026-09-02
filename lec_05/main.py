import os
import json
from pathlib import Path

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document


# ==========================================
# 1. SETUP
# ==========================================

# .env is in AI/.env
# main.py is in AI/lec_05/main.py

env_path = Path(__file__).resolve().parent.parent / ".env"
load_dotenv(env_path)

api_key = os.getenv("GROQ_API_KEY")

if not api_key:
    raise ValueError("GROQ_API_KEY not found")

client = Groq(api_key=api_key)

model = "openai/gpt-oss-120b"


# ==========================================
# 2. JOB SCHEMA
# ==========================================

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]


# ==========================================
# 3. CREATE JOB
# ==========================================

job = JobD(
    role="Frontend Developer Intern",

    required_skills=[
        "HTML",
        "CSS",
        "JavaScript",
        "ReactJS"
    ],

    preferred_skills=[
        "Git",
        "GitHub",
        "REST API",
        "Tailwind CSS"
    ],

    minimum_experience=0.0,

    education_requirements=[
        "B.Tech in Computer Science or related field"
    ],

    responsibilities=[
        "Build responsive web interfaces",
        "Develop reusable React components",
        "Work with APIs",
        "Collaborate with developers"
    ]
)


# ==========================================
# 4. RESUME SCHEMA
# ==========================================

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = Field(default_factory=list)
    experiences: list[Experience] = Field(default_factory=list)
    education: list[str] = Field(default_factory=list)
    projects: list[str] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)


# ==========================================
# 5. MATCH RESULT
# ==========================================

class MatchResult(BaseModel):
    score: float
    details: dict


# ==========================================
# 6. READ PDF
# ==========================================

def read_pdf(file_path):

    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:

        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


# ==========================================
# 7. READ DOCX
# ==========================================

def read_docx(file_path):

    document = Document(file_path)

    text = ""

    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():
                    text += cell.text + "\n"

    return text


# ==========================================
# 8. READ RESUME
# ==========================================

def read_resume(file_path):

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return read_pdf(file_path)

    elif extension == ".docx":
        return read_docx(file_path)

    return None


# ==========================================
# 9. PARSE RESUME USING LLM
# ==========================================

def parse_resume(resume_text):

    schema = Resume.model_json_schema()

    system_prompt = f"""
You are an expert resume parser.

Extract information from the resume.

Return ONLY valid JSON matching this schema:

{schema}

Rules:

1. Do not invent information.
2. Missing values should be null.
3. Missing lists should be empty.
4. Include internships inside experiences.
5. Extract skills from the entire resume.
"""

    user_prompt = f"""
Parse the following resume:

{resume_text}
"""

    messages = [

        {
            "role": "system",
            "content": system_prompt
        },

        {
            "role": "user",
            "content": user_prompt
        }

    ]

    response = client.chat.completions.create(

        model=model,

        messages=messages,

        response_format={
            "type": "json_object"
        }

    )

    raw_output = response.choices[0].message.content

    data = json.loads(raw_output)

    return Resume(**data)


# ==========================================
# 10. SCORE CANDIDATE
# ==========================================

def final_score(job, resume):

    schema = MatchResult.model_json_schema()

    prompt = f"""
You are an HR recruiter.

Compare this candidate's resume with the job description.

JOB DESCRIPTION:

{job.model_dump_json(indent=2)}

CANDIDATE RESUME:

{resume.model_dump_json(indent=2)}

Return JSON matching this schema:

{schema}

The "details" object should contain:

1. candidate_name
2. matching_skills
3. missing_important_skills
4. experience_requirement_met
5. final_verdict

The "score" must be an overall match percentage
between 0 and 100.

Keep the response concise.
"""

    response = client.chat.completions.create(

        model=model,

        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],

        response_format={
            "type": "json_object"
        }

    )

    data = json.loads(
        response.choices[0].message.content
    )

    return MatchResult(**data)


# ==========================================
# 11. MAIN PROGRAM
# ==========================================

resume_folder = Path(__file__).resolve().parent / "resumes"

all_results = []


for file_path in resume_folder.iterdir():

    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue

    print("\n========================================")
    print("Processing:", file_path.name)
    print("========================================")

    resume_text = read_resume(file_path)

    if not resume_text:
        print("Could not extract text.")
        continue

    parsed_resume = parse_resume(resume_text)

    result = final_score(job, parsed_resume)

    print("Candidate:", parsed_resume.name)
    print("Score:", result.score)

    print("Details:")
    print(result.details)

    all_results.append({

        "name": parsed_resume.name,

        "score": result.score,

        "details": result.details

    })


# ==========================================
# 12. SORT
# ==========================================

all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)


# ==========================================
# 13. TOP AND WORST
# ==========================================

top_2 = all_results[:2]

worst_2 = all_results[-2:]


# ==========================================
# 14. PRINT TOP 2
# ==========================================

print("\n\nTOP 2 CANDIDATES")
print("============================")

for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])


# ==========================================
# 15. PRINT LOWEST 2
# ==========================================

print("\n\nLOWEST 2 CANDIDATES")
print("============================")

for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])